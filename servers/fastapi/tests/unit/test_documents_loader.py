import asyncio
import json

import pytest
from fastapi import HTTPException

from services.documents_loader import (
    DocumentsLoader,
    _unwrap_liteparse_json_line_if_stored,
    clean_extracted_document_text,
)


def test_unwrap_liteparse_json_line_extracts_text_field():
    inner_text = "Title\n\nBody with \"quotes\""
    payload = json.dumps({"ok": True, "filePath": "/tmp/test.pdf", "text": inner_text})

    assert _unwrap_liteparse_json_line_if_stored(payload) == inner_text
    assert _unwrap_liteparse_json_line_if_stored(f"  {payload}") == inner_text


def test_unwrap_liteparse_json_line_leaves_non_json_text():
    plain_text = "Not JSON, should stay as-is."
    assert _unwrap_liteparse_json_line_if_stored(plain_text) == plain_text


def test_clean_extracted_document_text_handles_malformed_json_body():
    malformed = (
        '{"ok": true, "filePath": "/tmp/test.pdf", "text": '
        '"hello\\nworld\\u0021 and trailing'
    )
    cleaned = clean_extracted_document_text(malformed)
    assert cleaned == "hello\nworld! and trailing"


def test_clean_extracted_document_text_unwraps_nested_liteparse_payloads():
    nested = json.dumps(
        {
            "ok": True,
            "filePath": "/tmp/outer.pdf",
            "text": json.dumps(
                {"ok": True, "filePath": "/tmp/inner.pdf", "text": "final body"}
            ),
        }
    )
    assert clean_extracted_document_text(nested) == "final body"


def test_load_pdf_requires_temp_dir_when_images_are_requested():
    loader = DocumentsLoader(file_paths=[])

    with pytest.raises(HTTPException) as exc:
        asyncio.run(
            loader.load_pdf(
                file_path="/tmp/fake.pdf",
                load_text=False,
                load_images=True,
                temp_dir=None,
            )
        )

    assert exc.value.status_code == 400
    assert "temp_dir is required" in exc.value.detail


def test_docx_to_markdown_preserves_chinese_text_and_structure(tmp_path):
    from docx import Document
    from services.documents_loader import docx_to_markdown

    path = tmp_path / "中文方案.docx"
    document = Document()
    document.add_heading("人工智能项目方案", level=1)
    document.add_paragraph("这是第一段中文正文，包含关键背景和目标。")
    document.add_paragraph("第一项行动", style="List Bullet")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "阶段"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "启动"
    table.cell(1, 1).text = "完成需求确认"
    document.save(path)

    markdown = docx_to_markdown(str(path))

    assert "# 人工智能项目方案" in markdown
    assert "这是第一段中文正文，包含关键背景和目标。" in markdown
    assert "- 第一项行动" in markdown
    assert "阶段 | 说明" in markdown
    assert "启动 | 完成需求确认" in markdown
